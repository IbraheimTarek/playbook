import streamlit as st
import os
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain, create_history_aware_retriever
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_core.messages import AIMessage, HumanMessage
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables.history import RunnableWithMessageHistory
import pandas as pd
from langchain.docstore.document import Document
from dotenv import load_dotenv
import time
groq_api_key= os.getenv('GROQ_API_KEY')

llm = ChatGroq(groq_api_key = groq_api_key, model = 'Llama3-70b-8192')

### Contextualize question ###
contextualize_q_system_prompt = """Given a chat history and the latest user question \
which might reference context in the chat history, formulate a standalone question \
which can be understood without the chat history. Do NOT answer the question, \
just reformulate it if needed and otherwise return it as is."""

contextualize_q_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", contextualize_q_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ]
)

qa_system_prompt = """Answer this question using the provided context only and give the users all Actions and Recommendations needed to solve the problem. \
Use the following pieces of retrieved context to answer the question. \
If you don't know the answer, just say that you don't know. \

context:
{context}"""

qa_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", qa_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ]
)
def create_vector_embedding():
    if 'vectors' not in st.session_state:
        df = pd.read_excel('data.xlsx')
        st.session_state.documents = []
        for _, row in df.iterrows():
            content = row.to_string()  # Convert the row to a string
            doc = Document(page_content=content)
            st.session_state.documents.append(doc)
        st.session_state.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        st.session_state.vectors = FAISS.from_documents(st.session_state.documents,st.session_state.embeddings)

if 'store' not in st.session_state:
    st.session_state.store = {}
    
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'context' not in st.session_state:
    st.session_state.context = []

def get_session_history(session_id: str) -> BaseChatMessageHistory:
    if session_id not in st.session_state.store:
        st.session_state.store[session_id] = ChatMessageHistory()
    return st.session_state.store[session_id]


### streamlit ui ###

if "messages" not in st.session_state:
    st.session_state.messages = []

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])


user_prompt = st.chat_input("Enter Your Papers")

with st.sidebar:  # Alternatively, you can use st.container() for a specific area on the page
    if 'vectors' not in st.session_state:
        if st.button("Document Embedding"):
            create_vector_embedding()
            st.write("Vector Database is ready")
        else:
            st.write('Click the button first to initialize the chat vector store')
    else:
        st.write("Vector Database is loaded")

if user_prompt :
    retriever = st.session_state.vectors.as_retriever()
    question_answer_chain  = create_stuff_documents_chain(llm,qa_prompt)
    history_aware_retriever = create_history_aware_retriever(llm, retriever, contextualize_q_prompt)
    rag_chain  = create_retrieval_chain(history_aware_retriever,question_answer_chain)

    conversational_rag_chain = RunnableWithMessageHistory(
        rag_chain,
        get_session_history,
        input_messages_key="input",
        history_messages_key="chat_history",
        output_messages_key="answer",
    )

    response = conversational_rag_chain.invoke(
        {"input": user_prompt},
        config={
            "configurable": {"session_id": "abc123"}
        }, 
    )

    st.session_state.chat_history.append(HumanMessage(content=user_prompt))
    st.session_state.chat_history.append(AIMessage(content=response['answer']))

    model_response = response['answer']
    # st.write(model_response)

    st.session_state.messages.append({"role": "user", "content": user_prompt})

    with st.chat_message("user"):
        st.markdown(user_prompt)


    def chunk_response(text, chunk_size=20):
        for i in range(0, len(text), chunk_size):
            yield text[i:i + chunk_size]

    # Display assistant message with streaming effect
    with st.chat_message("assistant"):
        message_placeholder = st.empty()  # Placeholder to update the message content
        full_response = ""
        for chunk in chunk_response(model_response):
            full_response += chunk
            message_placeholder.markdown(full_response)
            time.sleep(0.1)  # Simulate a delay between chunks to mimic streaming
        with st.expander("Document similarity search"):
            for i, doc in enumerate(response['context']):
                st.write(doc.page_content)
                st.write("----------------------------")
                if doc.page_content not in st.session_state.context:
                    st.session_state.context.append(doc.page_content)


    st.session_state.messages.append({"role": "assistant", "content": model_response})

from langchain import PromptTemplate, LLMChain
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import re
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_text(text):
    cleaned = re.sub(r'\*+\s*', '', text)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def generate_incident_report(llm, chat_history, context):
    template = """
    Based on the following chat history and context, generate an incident report 
    with ONLY the following sections in this exact order:
    1. Executive Summary
    2. Scope of Work
    3. Incident Details
    4. Severity
    5. Impact
    6. Findings
    7. Action
    8. Recommendation
    9. Investigators

    For each section, provide relevant information without adding any sections or details that are not explicitly requested.
    For the Investigators section, list at least two names with their roles.

    Chat History: {chat_history}
    Context: {context}

    Please format the response as a structured document without any additional headers or information.
    """

    prompt = PromptTemplate(
        input_variables=["chat_history", "context"],
        template=template,
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    
    try:
        report_content = chain.run(chat_history=chat_history, context=context)
        logger.info(f"Generated report content: {report_content[:500]}...")  # Log first 500 characters
    except Exception as e:
        logger.error(f"Error generating report content: {str(e)}")
        return "Failed to generate report content."

    if not report_content.strip():
        logger.error("Generated report content is empty.")
        return "Generated report content is empty."

    doc = Document()

    styles = doc.styles
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)

    style_heading = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    font = style_heading.font
    font.name = 'Calibri'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(255, 0, 0)  # Red color

    doc.add_heading('Incident Report', 0)

    # Split the content into sections
    sections = report_content.split('\n\n')
    
    for section in sections:
        # Check if the section starts with a number (indicating a new section)
        match = re.match(r'^(\d+\.\s*)(.*?)(:?)\s*(.*)', section, re.DOTALL)
        if match:
            number, title, colon, content = match.groups()
            
            # Add the section title
            para = doc.add_paragraph(style='CustomHeading')
            para.add_run(f"{number}{title}{colon}")
            
            # Add the section content
            if content:
                doc.add_paragraph(clean_text(content), style='Normal')
        else:
            # If it's not a numbered section, just add it as normal text
            doc.add_paragraph(clean_text(section), style='Normal')

    try:
        doc.save('generated_incident_report.docx')
        logger.info("Document saved successfully.")
    except Exception as e:
        logger.error(f"Error saving document: {str(e)}")
        return "Failed to save document."

    return "generated_incident_report.docx"

# with st.sidebar:
#     st.write(st.session_state.context)
#     if st.button("Generate Report"):
#         result = generate_incident_report(llm, st.session_state.chat_history, st.session_state.context)
#         st.success(f"Report generated: {result}")
#         st.download_button('Download Report', open(result, 'rb'), file_name='incident_report.docx')
        
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False
if 'report_content' not in st.session_state:
    st.session_state.report_content = None

def generate_report():
    result = generate_incident_report(llm, st.session_state.chat_history, st.session_state.context)
    if result.endswith('.docx'):
        st.session_state.report_generated = True
        # Read the generated report into memory
        with open(result, 'rb') as file:
            st.session_state.report_content = file.read()
    else:
        st.error(f"Failed to generate report: {result}")

with st.sidebar:
    st.write(st.session_state.context)
    if st.button("Generate Report"):
        generate_report()

    if st.session_state.report_generated:
        st.success("Report generated successfully!")
        
        # Create a download button that doesn't trigger a page refresh
        st.download_button(
            label="Download Report",
            data=st.session_state.report_content,
            file_name="incident_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
## test
### the CEO Received a fraudulent email impersonating him, requesting immediate transfer of funds. also some of employees that they have unidentified malware infected systems.###